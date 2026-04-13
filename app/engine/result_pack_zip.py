"""Build the v0.1 ZIP export with only useful generated files."""

from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path
from typing import Dict, Optional

from docx import Document

BAD_TOKENS = (
    "未提供",
    "请参见以下",
    "这里可以补充",
    "你可以这样回答",
    "某公司",
    "某项目",
    "xx",
    "XX",
    "××",
)


def clean_generated_text(body: str, *, resume: bool = False) -> str:
    text = body or ""
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"<!--[\s\S]*?-->", "", text)
    lines = []
    bad_resume_phrases = ("建议补充", "可进一步说明", "可以考虑", "这里建议", "应当强调")
    internal = (
        "general",
        "education",
        "projects",
        "accomplishments",
        "Certifications",
        "programming languages",
        "suggestion id",
    )
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            lines.append("")
            continue
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"^>\s*", "", line)
        line = re.sub(r"^\s*[-*]\s+", "", line)
        if any(tok in line for tok in BAD_TOKENS):
            continue
        if any(tok in line for tok in internal):
            continue
        if resume and any(p in line for p in bad_resume_phrases):
            continue
        lines.append(line)
    out = "\n".join(lines)
    out = re.sub(r"\n{3,}", "\n\n", out).strip()
    return out


def is_meaningful_generated_text(body: str, *, min_chars: int = 40) -> bool:
    text = clean_generated_text(body)
    compact = re.sub(r"\s+", "", text)
    if len(compact) < min_chars:
        return False
    if any(tok in compact for tok in BAD_TOKENS):
        return False
    return True


def _text_to_docx_bytes(title: str, body: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for block in (body or "").split("\n\n"):
        b = block.strip()
        if not b:
            continue
        for line in b.split("\n"):
            clean = line.strip()
            if clean:
                doc.add_paragraph(clean)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_result_pack_zip_bytes(
    *,
    resume_markdown: str,
    job_brief_md: str,
    resume_delta_md: str,
    interview_pack_md: str,
    interview_report_md: Optional[str] = None,
    include_interview_report: bool = False,
    capabilities: Optional[Dict[str, object]] = None,
) -> bytes:
    caps = capabilities or {}
    resume_body = clean_generated_text(resume_markdown, resume=True)
    job_body = clean_generated_text(job_brief_md)
    delta_body = clean_generated_text(resume_delta_md)
    interview_body = clean_generated_text(interview_pack_md)
    report_body = clean_generated_text(interview_report_md or "")

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        if caps.get("can_generate_resume", True) and is_meaningful_generated_text(resume_body, min_chars=60):
            zf.writestr("主稿-简历.docx", _text_to_docx_bytes("主稿-简历", resume_body))
        if caps.get("can_generate_resume_suggestions", True) and is_meaningful_generated_text(delta_body, min_chars=60):
            zf.writestr("简历修改建议.docx", _text_to_docx_bytes("简历修改建议", delta_body))
        if caps.get("can_generate_job_brief") and is_meaningful_generated_text(job_body, min_chars=60):
            zf.writestr("岗位理解卡.docx", _text_to_docx_bytes("岗位理解卡", job_body))
        if caps.get("can_generate_interview_prep") and is_meaningful_generated_text(interview_body, min_chars=60):
            zf.writestr("面试准备包.docx", _text_to_docx_bytes("面试准备包", interview_body))
        if (
            include_interview_report
            and caps.get("can_generate_report")
            and is_meaningful_generated_text(report_body, min_chars=80)
        ):
            zf.writestr("面试练习综合分析报告.docx", _text_to_docx_bytes("面试练习综合分析报告", report_body))
    return buf.getvalue()


def read_generated_md(generated_dir: Path, name: str) -> str:
    p = generated_dir / name
    if not p.is_file():
        return ""
    try:
        return p.read_text(encoding="utf-8")
    except Exception:
        return ""
