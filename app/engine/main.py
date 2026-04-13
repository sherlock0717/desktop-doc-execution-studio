from pathlib import Path
from datetime import datetime
import json
import os
import re
import shutil
import uuid
from typing import Any, Dict, List, Optional, Set
from urllib.parse import quote

from fastapi import FastAPI, HTTPException, UploadFile, File, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from docx import Document
from pypdf import PdfReader
from pydantic import BaseModel, Field, ValidationError

from .coaching_generator import generate_interview_coaching_markdown
from .demo_payloads import build_demo_payloads
from .execution_engine import (
    apply_goal_mode_truncation,
    build_comparison_summary,
    case_runtime_dir,
    ensure_interview_talking_draft,
    load_execution_state,
    merge_suggestion_status,
    new_generation_id,
    resume_payload_to_suggestion_items,
    save_execution_state,
    working_draft_paths,
    working_draft_status,
)
from .readiness import build_readiness_summary, build_upload_snapshot
from .refined_draft import (
    build_accepted_applied_summary,
    build_compare_summary,
    hash_text,
    load_latest_refined,
    load_refined_by_id,
    messages_refined_resume,
    new_refined_id,
    preview_for_execution,
    save_refined_artifacts,
)
from .interview_practice import (
    generate_interview_practice_pack_safe,
    load_practice_pack,
    save_practice_pack,
    score_interview_answer,
    submit_practice_full_report_safe,
)
from .material_validation import validate_materials_for_parse
from .result_pack_zip import build_result_pack_zip_bytes, read_generated_md
from .resume_draft_builder import (
    SECTION_TITLES,
    build_resume_section_views,
    ensure_structured_resume_draft,
    get_section_body_text,
    insert_improvement_block_near_section_header,
    load_original_resume_snapshot,
    migrate_old_template_to_clean,
    normalize_target_section,
    parse_resume_draft_sections,
    remove_improvement_block,
    replace_section_body_in_file,
    save_draft_apply_summary,
    save_original_resume_snapshot,
    strip_improvement_blocks_for_llm,
)
from .suggestion_rewrite import rewrite_resume_section_with_llm
from .llm_provider import (
    check_ollama_available,
    extract_json_object,
    get_active_model_label,
    get_provider,
)
from .markdown_renderers import (
    render_interview_pack_md,
    render_job_brief_md,
    render_resume_delta_md,
)
from .product_constants import PROMPT_VERSION
from .prompt_builders import (
    build_import_context,
    messages_interview_pack,
    messages_job_brief,
    messages_resume_delta,
)
from .schema_definitions import (
    InterviewPackPayload,
    InterviewPracticePayload,
    JobBriefPayload,
    RefinedResumePayload,
    ResumeDeltaPayload,
)
from .snippet_extractor import SnippetRecord, build_snippet_catalog, catalog_to_map

app = FastAPI(title="Desktop Job Materials Studio API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

ROOT_DIR = Path(__file__).resolve().parents[2]
INPUT_DIR = ROOT_DIR / "cases" / "sample_case_01" / "input"
DEMO_OUTPUT_DIR = ROOT_DIR / "outputs" / "demo_outputs" / "sample_case_01"
UPLOAD_CASE_DIR = ROOT_DIR / "uploads" / "current_case"
GENERATED_OUTPUT_DIR = ROOT_DIR / "outputs" / "generated" / "current_case"

OUTPUT_FILE_MAP = {
    "job_brief": ("岗位理解卡", "job_brief.md", "markdown"),
    "resume_delta": ("简历修改建议", "resume_delta.md", "markdown"),
    "interview_pack": ("面试准备包", "interview_pack.md", "markdown"),
    "archive_manifest": ("解析摘要", "archive_manifest.json", "json"),
}

DOC_LABELS = {
    "jd": "岗位说明",
    "resume": "简历",
    "interview_note": "面试记录",
    "supporting_material": "补充材料",
}

DEMO_DOCUMENTS = [
    {"name": "xiaohe_jd_pm.txt", "type": "jd", "label": "岗位说明"},
    {"name": "xiaohe_resume.txt", "type": "resume", "label": "简历"},
    {"name": "xiaohe_interview_notes.md", "type": "interview_note", "label": "面试记录"},
    {"name": "xiaohe_project_notes.md", "type": "supporting_material", "label": "项目补充"},
]


class CasePatch(BaseModel):
    title: Optional[str] = None
    goal_mode: Optional[str] = None


class DocTypePatch(BaseModel):
    doc_type: str = Field(..., description="jd | resume | interview_note | supporting_material")


class SuggestionStatusPatch(BaseModel):
    status: str = Field(..., description="仅支持 accepted")


class WorkingDraftPut(BaseModel):
    content: str = Field(..., description="Markdown body")


class InterviewScoreBody(BaseModel):
    question: str = Field(..., description="题目全文")
    user_answer: str = Field("", description="用户作答")
    reference_answer: str = Field("", description="参考回答，用于对齐评分标准")


class ResultPackExportBody(BaseModel):
    include_interview_report: bool = Field(False, description="是否打包面试综合分析报告")


class HrPersonaBody(BaseModel):
    persona_id: str = Field(..., description="严厉专业型等 stable id")
    label: Optional[str] = Field(None, description="展示用中文标签")


def ensure_dir(path: Path):
    path.mkdir(parents=True, exist_ok=True)


def clear_directory(path: Path):
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def read_text_file(path: Path) -> str:
    if not path.exists():
        return ""
    for encoding in ["utf-8-sig", "utf-8", "gbk"]:
        try:
            return path.read_text(encoding=encoding)
        except Exception:
            continue
    return ""


def sanitize_zip_base(name: str) -> str:
    s = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", (name or "").strip())
    return (s[:80] or "case").strip(" .") or "case"


def build_materials_blob(documents: List[Dict], text_map: Dict[str, str], limit: int = 28000) -> str:
    parts: List[str] = []
    for d in documents or []:
        nm = d.get("name", "")
        lab = d.get("label", "")
        body = (text_map or {}).get(nm, "") or ""
        parts.append(f"【{lab}|{nm}】\n{body[:12000]}")
    blob = "\n\n---\n\n".join(parts)
    return blob[:limit]


def hr_persona_path(rt: Path) -> Path:
    return rt / "hr_persona.json"


def default_hr_persona() -> Dict[str, Any]:
    return {"persona_id": "calm_rational", "locked": False, "label": "冷静理性型"}


def load_hr_persona(rt: Path) -> Dict[str, Any]:
    p = hr_persona_path(rt)
    if not p.is_file():
        return default_hr_persona()
    try:
        data = json.loads(p.read_text(encoding="utf-8"))
        if isinstance(data, dict):
            merged = default_hr_persona()
            merged.update(data)
            return merged
    except Exception:
        pass
    return default_hr_persona()


def read_docx_file(path: Path) -> str:
    try:
        doc = Document(path)
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append(" | ".join(row_texts))

        return "\n".join(parts)
    except Exception:
        return ""


def read_pdf_file(path: Path) -> str:
    try:
        reader = PdfReader(str(path))
        parts = []
        for page in reader.pages[:20]:
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                parts.append(text)
        return "\n".join(parts)
    except Exception:
        return ""


def write_text_file(path: Path, content: str):
    ensure_dir(path.parent)
    path.write_text(content, encoding="utf-8")


def extract_text_for_generation(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix in [".txt", ".md", ".json"]:
        return read_text_file(path)
    if suffix == ".docx":
        return read_docx_file(path)
    if suffix == ".pdf":
        return read_pdf_file(path)

    return ""


def get_first_meaningful_lines(text: str, max_lines: int = 2) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return "当前文件已导入，但还未抽取到可用正文。"
    return " ".join(lines[:max_lines])[:180]


def classify_document(filename: str, text: str):
    name = filename.lower()
    text_head = text[:2500]

    scores = {
        "jd": 0,
        "resume": 0,
        "interview_note": 0,
        "supporting_material": 0,
    }

    jd_name_keywords = ["jd", "job", "职位", "岗位说明", "岗位需求", "招聘", "招聘信息"]
    resume_name_keywords = ["resume", "cv", "简历", "个人简历"]
    interview_name_keywords = ["interview", "面试", "访谈", "复盘", "记录"]
    support_name_keywords = ["说明", "notes", "memo", "补充"]

    jd_text_keywords = ["岗位职责", "岗位描述", "任职要求", "职位描述", "职位要求", "岗位名称", "工作地点", "薪资范围"]
    resume_text_keywords = ["姓名", "教育背景", "项目经历", "实习经历", "校园经历", "技能", "自我评价", "联系方式", "求职意向"]
    interview_text_keywords = ["面试官", "追问", "一面", "二面", "hr面", "面试反馈", "复盘", "问题记录"]
    support_text_keywords = ["补充说明", "备注", "说明", "项目目标", "背景"]

    for keyword in jd_name_keywords:
        if keyword in name:
            scores["jd"] += 3

    for keyword in resume_name_keywords:
        if keyword in name:
            scores["resume"] += 4

    for keyword in interview_name_keywords:
        if keyword in name:
            scores["interview_note"] += 4

    for keyword in support_name_keywords:
        if keyword in name:
            scores["supporting_material"] += 2

    for keyword in jd_text_keywords:
        if keyword in text_head:
            scores["jd"] += 2

    for keyword in resume_text_keywords:
        if keyword in text_head:
            scores["resume"] += 3

    for keyword in interview_text_keywords:
        if keyword in text_head:
            scores["interview_note"] += 3

    for keyword in support_text_keywords:
        if keyword in text_head:
            scores["supporting_material"] += 1

    best_type = max(scores, key=scores.get)

    if scores[best_type] <= 0:
        return "supporting_material", "补充材料"

    labels = {
        "jd": "岗位说明",
        "resume": "简历",
        "interview_note": "面试记录",
        "supporting_material": "补充材料",
    }

    return best_type, labels[best_type]


def compute_missing_inputs(goal_mode: str, documents: List[Dict]) -> List[str]:
    types = {d["type"] for d in documents}
    missing: List[str] = []
    if goal_mode in ("delivery", "both"):
        if "jd" not in types:
            missing.append("岗位说明（JD）")
        if "resume" not in types:
            missing.append("简历")
    if goal_mode in ("interview", "both"):
        if "interview_note" not in types:
            missing.append("面试记录 / 复盘")
    if goal_mode == "both" and "supporting_material" not in types:
        missing.append("补充材料（可选，但建议提供）")
    return missing


def compute_acceptance_hints(
    documents: List[Dict],
    catalog: List[SnippetRecord],
    goal_mode: str,
    missing_inputs: List[str],
) -> Dict[str, Any]:
    n = len(catalog)
    unconfirmed = sum(1 for d in documents if not d.get("type_confirmed"))
    avg_len = 0.0
    if n:
        avg_len = sum(len(c["snippet_text"]) for c in catalog) / n
    weak_evidence = n < 3 or (n > 0 and avg_len < 45)
    material_insufficient = len(missing_inputs) > 0
    needs_human_review = material_insufficient or weak_evidence or unconfirmed > 0
    return {
        "material_insufficient": material_insufficient,
        "weak_evidence": weak_evidence,
        "needs_human_review": needs_human_review,
        "unconfirmed_type_count": unconfirmed,
    }


def derive_input_capabilities(documents: List[Dict]) -> Dict[str, Any]:
    types = {d.get("type") for d in documents or []}
    unconfirmed = [d.get("name") for d in documents or [] if not d.get("type_confirmed")]
    has_resume = "resume" in types
    has_jd = "jd" in types
    has_interview_note = "interview_note" in types
    has_support = "supporting_material" in types
    can_resume = has_resume
    can_job = has_resume and has_jd
    can_interview_prep = has_resume and has_jd
    can_practice = has_resume and has_jd and (has_interview_note or has_support)
    caps = {
        "can_generate_resume": can_resume,
        "can_generate_resume_suggestions": can_resume,
        "can_generate_job_brief": can_job,
        "can_generate_interview_prep": can_interview_prep,
        "can_generate_practice": can_practice,
        "can_generate_report": can_practice,
        "unconfirmed_type_count": len(unconfirmed),
        "unconfirmed_files": unconfirmed,
        "detected_types": sorted(t for t in types if t),
    }
    if unconfirmed:
        msg = "请先确认全部材料类型，再开始解析。"
    elif can_practice:
        msg = "当前已检测到简历、岗位说明及补充/面试材料。可完整支持简历优化、岗位理解、面试辅导与练习报告。"
    elif can_job:
        msg = "当前已检测到简历与岗位说明。可进行简历优化、岗位理解与基础面试辅导。"
    elif can_resume:
        msg = "当前仅检测到简历材料。可进行简历优化；岗位理解、面试准备与练习相关内容将不会生成。"
    else:
        msg = "当前材料不足以解析。请至少导入并确认一份简历。"
    caps["summary"] = msg
    return caps


def output_ids_for_capabilities(caps: Dict[str, Any], *, include_manifest: bool = True) -> List[str]:
    ids: List[str] = []
    if caps.get("can_generate_job_brief"):
        ids.append("job_brief")
    if caps.get("can_generate_resume_suggestions"):
        ids.append("resume_delta")
    if caps.get("can_generate_interview_prep"):
        ids.append("interview_pack")
    if include_manifest:
        ids.append("archive_manifest")
    return ids


def build_output_list(status: str, capabilities: Optional[Dict[str, Any]] = None):
    outputs = []
    ids = output_ids_for_capabilities(capabilities or {}, include_manifest=True) if capabilities else list(OUTPUT_FILE_MAP.keys())
    for output_id in ids:
        title, filename, output_format = OUTPUT_FILE_MAP[output_id]
        outputs.append(
            {
                "id": output_id,
                "name": title,
                "status": status,
                "format": output_format,
            }
        )
    return outputs


def collect_uploaded_texts(documents: List[Dict]):
    text_map: Dict[str, str] = {}
    for doc in documents:
        file_path = UPLOAD_CASE_DIR / doc["name"]
        text_map[doc["name"]] = extract_text_for_generation(file_path)
    return text_map


def collect_typed_text(documents: List[Dict], text_map: Dict[str, str], doc_type: str) -> str:
    parts: List[str] = []
    for doc in documents:
        if doc.get("type") == doc_type:
            parts.append(text_map.get(doc["name"], ""))
    return "\n\n".join(parts).strip()


def write_interview_expression_coaching(runtime_dir: Path, documents: List[Dict], text_map: Dict[str, str]) -> None:
    """面试表达辅导：优先 LLM 个性化批注；模型不可用时回落为静态摘录。"""
    persona = load_hr_persona(runtime_dir)
    blob = build_materials_blob(documents, text_map)
    p = runtime_dir / "interview_expression_coaching.md"
    runtime_dir.mkdir(parents=True, exist_ok=True)
    backend = os.getenv("LLM_BACKEND", "ollama").lower().strip()
    if backend in ("ollama", ""):
        use_llm = check_ollama_available()
    else:
        use_llm = True
    if use_llm:
        try:
            body = generate_interview_coaching_markdown(
                materials_blob=blob,
                persona_id=str(persona.get("persona_id") or "calm_rational"),
            )
            p.write_text(body, encoding="utf-8")
            return
        except Exception:
            pass
    resume = collect_typed_text(documents, text_map, "resume")
    jd = collect_typed_text(documents, text_map, "jd")
    sup = collect_typed_text(documents, text_map, "supporting_material")
    body = (
        "面试表达辅导\n\n"
        "模型暂不可用，以下为材料摘录供你对照练习。连接模型后可重新解析或稍后在本页触发更新。\n\n"
        "简历摘录：\n"
        f"{resume[:4500] or '（暂无）'}\n\n"
        "岗位摘录：\n"
        f"{jd[:4500] or '（暂无）'}\n\n"
        "补充材料：\n"
        f"{sup[:2500] or '（暂无）'}\n"
    )
    p.write_text(body, encoding="utf-8")


def build_pending_accept_overlays(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for s in items or []:
        if s.get("status") == "accepted" and not s.get("applied_to_draft"):
            out.append(
                {
                    "id": s.get("id"),
                    "target_section": s.get("target_section"),
                    "text": (s.get("text") or "").strip(),
                }
            )
    return out


def enrich_execution(case_id: str) -> Dict[str, Any]:
    rt = case_runtime_dir(ROOT_DIR, case_id)
    ex = load_execution_state(rt / "execution.json")
    rp, ip = working_draft_paths(rt)
    summary_path = rt / "draft_apply_summary.json"
    draft_apply_summary = None
    if summary_path.exists():
        try:
            draft_apply_summary = json.loads(summary_path.read_text(encoding="utf-8"))
        except Exception:
            draft_apply_summary = None
    if not ex:
        return {
            "generation_id": None,
            "previous_generation_id": None,
            "suggestion_items": [],
            "comparison_summary": None,
            "input_files": [],
            "missing_inputs": [],
            "working_drafts": working_draft_status(rp, ip),
            "draft_apply_summary": draft_apply_summary,
            "resume_sections": [],
            "pending_accept_overlays": [],
            "workflow_phase": "none",
            "refined_draft_preview": None,
            "hr_persona": load_hr_persona(rt),
        }
    merged = {**ex}
    merged["working_drafts"] = working_draft_status(rp, ip)
    merged["draft_apply_summary"] = draft_apply_summary
    merged["hr_persona"] = load_hr_persona(rt)
    rtext = read_text_file(rp) if rp.exists() else ""
    orig = load_original_resume_snapshot(rt)
    items = merged.get("suggestion_items") or []
    merged["resume_sections"] = build_resume_section_views(rtext, items, orig)
    merged["pending_accept_overlays"] = build_pending_accept_overlays(items)
    merged.setdefault("workflow_phase", "parsed" if merged.get("generation_id") else "none")
    merged["refined_draft_preview"] = preview_for_execution(rt)
    return merged


def refresh_derived_case_fields():
    global STATE
    if STATE.get("mode") != "uploaded":
        return
    docs = STATE.get("documents") or []
    gm = STATE.get("goal_mode", "both")
    if not docs:
        STATE["missing_inputs"] = compute_missing_inputs(gm, [])
        STATE["snippet_count"] = 0
        STATE["input_capabilities"] = derive_input_capabilities([])
        STATE["outputs"] = build_output_list("pending", STATE["input_capabilities"])
        STATE["acceptance_hints"] = compute_acceptance_hints([], [], gm, STATE["missing_inputs"])
        if STATE.get("case_id"):
            STATE["execution"] = enrich_execution(STATE["case_id"])
        return
    text_map = collect_uploaded_texts(docs)
    catalog = build_snippet_catalog(docs, text_map)
    STATE["snippet_count"] = len(catalog)
    STATE["missing_inputs"] = compute_missing_inputs(gm, docs)
    STATE["input_capabilities"] = derive_input_capabilities(docs)
    if STATE.get("workflow_phase") != "parsed":
        STATE["outputs"] = build_output_list("pending", STATE["input_capabilities"])
    STATE["acceptance_hints"] = compute_acceptance_hints(
        docs,
        catalog,
        gm,
        STATE["missing_inputs"],
    )
    if STATE.get("case_id"):
        STATE["execution"] = enrich_execution(STATE["case_id"])


def finalize_cited_ids(ids: List[str], valid: Set[str], catalog: List[SnippetRecord]) -> List[str]:
    out = [i for i in ids if i in valid][:24]
    if not out and catalog:
        return [catalog[0]["snippet_id"]]
    return out


def resolve_cited_snippets_to_evidence(ids: List[str], cmap: Dict[str, SnippetRecord]) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for sid in ids:
        rec = cmap.get(sid)
        if not rec:
            continue
        out.append(
            {
                "snippet_id": rec["snippet_id"],
                "source_file": rec["source_file"],
                "snippet": rec["snippet_text"][:2000],
                "confidence": "中",
            }
        )
    return out


def _sanitize_payload(
    payload: JobBriefPayload | ResumeDeltaPayload | InterviewPackPayload,
    valid: Set[str],
    catalog: List[SnippetRecord],
):
    ids = finalize_cited_ids(list(payload.cited_snippet_ids), valid, catalog)
    return payload.model_copy(update={"cited_snippet_ids": ids})


def _run_llm_json_pack(messages, model_cls):
    provider = get_provider()
    raw = provider.complete(messages, json_mode=True)
    data = extract_json_object(raw)
    return model_cls.model_validate(data)


def build_evidence_item(doc: Optional[Dict], text_map: Dict, confidence: str = "中"):
    if not doc:
        return None

    snippet = get_first_meaningful_lines(text_map.get(doc["name"], ""))
    return {
        "snippet_id": "",
        "source_file": doc["name"],
        "snippet": snippet,
        "confidence": confidence,
    }


def _generation_meta_block(model_name: str, generated_at: str) -> Dict[str, str]:
    return {
        "model_name": model_name,
        "prompt_version": PROMPT_VERSION,
        "generated_at": generated_at,
    }


def write_generated_outputs(
    out_dir: Path,
    job_payload: JobBriefPayload,
    resume_payload: ResumeDeltaPayload,
    interview_payload: InterviewPackPayload,
    archive_manifest: Dict[str, Any],
):
    job_brief_content = render_job_brief_md(job_payload)
    resume_delta_content = render_resume_delta_md(resume_payload)
    interview_pack_content = render_interview_pack_md(interview_payload)

    write_text_file(out_dir / "job_brief.md", job_brief_content)
    write_text_file(out_dir / "resume_delta.md", resume_delta_content)
    write_text_file(out_dir / "interview_pack.md", interview_pack_content)
    write_text_file(
        out_dir / "job_brief.payload.json",
        json.dumps(job_payload.model_dump(mode="json"), ensure_ascii=False, indent=2),
    )
    write_text_file(
        out_dir / "resume_delta.payload.json",
        json.dumps(resume_payload.model_dump(mode="json"), ensure_ascii=False, indent=2),
    )
    write_text_file(
        out_dir / "interview_pack.payload.json",
        json.dumps(interview_payload.model_dump(mode="json"), ensure_ascii=False, indent=2),
    )
    write_text_file(
        out_dir / "archive_manifest.json",
        json.dumps(archive_manifest, ensure_ascii=False, indent=2),
    )


def build_output_definitions(
    out_dir: Path,
    cmap: Dict[str, SnippetRecord],
    catalog: List[SnippetRecord],
    job_payload: JobBriefPayload,
    resume_payload: ResumeDeltaPayload,
    interview_payload: InterviewPackPayload,
    generation_meta: Dict[str, str],
    acceptance: Dict[str, Any],
    capabilities: Optional[Dict[str, Any]] = None,
):
    job_ids = finalize_cited_ids(list(job_payload.cited_snippet_ids), set(cmap.keys()), catalog)
    resume_ids = finalize_cited_ids(list(resume_payload.cited_snippet_ids), set(cmap.keys()), catalog)
    interview_ids = finalize_cited_ids(list(interview_payload.cited_snippet_ids), set(cmap.keys()), catalog)

    archive_evidence: List[Dict[str, Any]] = []
    if catalog:
        first = catalog[0]
        archive_evidence.append(
            {
                "snippet_id": first["snippet_id"],
                "source_file": first["source_file"],
                "snippet": first["snippet_text"][:2000],
                "confidence": "低",
            }
        )

    definitions = {
        "job_brief": {
            "title": "岗位理解卡",
            "filename": "job_brief.md",
            "format": "markdown",
            "file_path": str(out_dir / "job_brief.md"),
            "evidence": resolve_cited_snippets_to_evidence(job_ids, cmap),
            "generation_meta": generation_meta,
            "acceptance": acceptance,
        },
        "resume_delta": {
            "title": "简历修改建议",
            "filename": "resume_delta.md",
            "format": "markdown",
            "file_path": str(out_dir / "resume_delta.md"),
            "evidence": resolve_cited_snippets_to_evidence(resume_ids, cmap),
            "generation_meta": generation_meta,
            "acceptance": acceptance,
        },
        "interview_pack": {
            "title": "面试准备包",
            "filename": "interview_pack.md",
            "format": "markdown",
            "file_path": str(out_dir / "interview_pack.md"),
            "evidence": resolve_cited_snippets_to_evidence(interview_ids, cmap),
            "generation_meta": generation_meta,
            "acceptance": acceptance,
        },
        "archive_manifest": {
            "title": "解析摘要",
            "filename": "archive_manifest.json",
            "format": "json",
            "file_path": str(out_dir / "archive_manifest.json"),
            "evidence": archive_evidence,
            "generation_meta": generation_meta,
            "acceptance": acceptance,
        },
    }
    if capabilities:
        allowed = set(output_ids_for_capabilities(capabilities, include_manifest=True))
        definitions = {k: v for k, v in definitions.items() if k in allowed}
    return definitions


def build_demo_state():
    documents: List[Dict] = []
    text_map: Dict[str, str] = {}
    for item in DEMO_DOCUMENTS:
        file_path = INPUT_DIR / item["name"]
        documents.append(
            {
                "name": item["name"],
                "type": item["type"],
                "label": item["label"],
                "exists": file_path.exists(),
                "extractable": True,
                "type_confirmed": True,
                "auto_type": item["type"],
            }
        )
        text_map[item["name"]] = extract_text_for_generation(file_path) if file_path.exists() else ""

    catalog = build_snippet_catalog(documents, text_map)
    cmap = catalog_to_map(catalog)
    job_p, resume_p, interview_p = build_demo_payloads(catalog)

    job_p = _sanitize_payload(job_p, set(cmap.keys()), catalog)
    resume_p = _sanitize_payload(resume_p, set(cmap.keys()), catalog)
    interview_p = _sanitize_payload(interview_p, set(cmap.keys()), catalog)

    now = datetime.now().isoformat(timespec="seconds")
    demo_meta = _generation_meta_block("demo_static", now)
    missing = compute_missing_inputs("both", documents)
    acceptance = compute_acceptance_hints(documents, catalog, "both", missing)
    capabilities = derive_input_capabilities(documents)

    ensure_dir(DEMO_OUTPUT_DIR)
    archive_manifest = {
        "case_id": "case_demo_001",
        "case_title": "示例：小何 · 产品经理岗位材料包",
        "generated_at": now,
        "input_files": [
            {
                "name": doc["name"],
                "type": doc["type"],
                "label": doc["label"],
                "extractable": doc["extractable"],
            }
            for doc in documents
        ],
        "generated_outputs": [
            "job_brief.md",
            "job_brief.payload.json",
            "resume_delta.md",
            "resume_delta.payload.json",
            "interview_pack.md",
            "interview_pack.payload.json",
            "archive_manifest.json",
        ],
        "notes": "演示案例：与真实模式共用 JSON→Markdown 渲染链；证据由 snippet_id 映射；解析摘要为确定性生成。",
    }

    write_generated_outputs(DEMO_OUTPUT_DIR, job_p, resume_p, interview_p, archive_manifest)

    rt_demo = case_runtime_dir(ROOT_DIR, "case_demo_001")
    demo_gen_id = "gen_demo_static"
    demo_items = resume_payload_to_suggestion_items(resume_p, demo_gen_id)
    save_execution_state(
        rt_demo / "execution.json",
        {
            "generation_id": demo_gen_id,
            "previous_generation_id": None,
            "input_files": [d["name"] for d in documents],
            "missing_inputs": missing,
            "suggestion_items": demo_items,
            "comparison_summary": None,
            "last_updated": now,
            "workflow_phase": "parsed",
            "input_snapshot_at_generation": build_upload_snapshot(INPUT_DIR, [d["name"] for d in documents]),
        },
    )
    save_draft_apply_summary(rt_demo, demo_items)
    resume_body_demo = collect_typed_text(documents, text_map, "resume")
    interview_body_demo = collect_typed_text(documents, text_map, "interview_note")
    drp, dip = working_draft_paths(rt_demo)
    ensure_structured_resume_draft(drp, resume_body_demo)
    ensure_interview_talking_draft(dip, interview_body_demo)
    save_original_resume_snapshot(rt_demo, read_text_file(drp))
    write_interview_expression_coaching(rt_demo, documents, text_map)

    definitions = build_output_definitions(
        DEMO_OUTPUT_DIR,
        cmap,
        catalog,
        job_p,
        resume_p,
        interview_p,
        demo_meta,
        acceptance,
        capabilities,
    )

    outputs = []
    for output_id in output_ids_for_capabilities(capabilities, include_manifest=True):
        definition = definitions.get(output_id)
        if not definition:
            continue
        outputs.append(
            {
                "id": output_id,
                "name": OUTPUT_FILE_MAP[output_id][0],
                "status": "ready",
                "format": OUTPUT_FILE_MAP[output_id][2],
            }
        )

    return {
        "mode": "demo",
        "case_id": "case_demo_001",
        "title": "示例：小何 · 产品经理岗位材料包",
        "goal_mode": "both",
        "documents": documents,
        "outputs": outputs,
        "stats": {
            "document_count": len(documents),
            "output_count": len(outputs),
        },
        "output_definitions": definitions,
        "missing_inputs": missing,
        "snippet_count": len(catalog),
        "acceptance_hints": acceptance,
        "input_capabilities": capabilities,
        "last_generated_at": now,
        "workflow_phase": "parsed",
        "execution": enrich_execution("case_demo_001"),
    }


# 初始状态为空白态，只有在用户主动操作时才加载案例
STATE: Dict[str, Any] = {
    "mode": "blank",
    "case_id": None,
    "title": "",
    "goal_mode": "both",
    "documents": [],
    "outputs": build_output_list("pending", derive_input_capabilities([])),
    "stats": {
        "document_count": 0,
        "output_count": len(OUTPUT_FILE_MAP),
    },
    "output_definitions": {},
    "missing_inputs": [],
    "snippet_count": 0,
    "acceptance_hints": None,
    "input_capabilities": derive_input_capabilities([]),
    "last_generated_at": None,
    "workflow_phase": "none",
    "execution": None,
}


def build_uploaded_state(
    documents: List[Dict],
    *,
    case_id: Optional[str] = None,
    title: str = "当前求职案例",
    goal_mode: str = "both",
) -> Dict[str, Any]:
    return {
        "mode": "uploaded",
        "case_id": case_id or f"case_{uuid.uuid4().hex[:10]}",
        "title": title,
        "goal_mode": goal_mode,
        "documents": documents,
        "outputs": build_output_list("pending", derive_input_capabilities(documents)),
        "stats": {
            "document_count": len(documents),
            "output_count": len(OUTPUT_FILE_MAP),
        },
        "output_definitions": {},
        "missing_inputs": [],
        "snippet_count": 0,
        "acceptance_hints": None,
        "input_capabilities": derive_input_capabilities(documents),
        "last_generated_at": None,
        "workflow_phase": "none",
    }


def save_uploaded_files(files: List[UploadFile]):
    clear_directory(UPLOAD_CASE_DIR)

    documents = []
    for index, file in enumerate(files, start=1):
        safe_name = Path(file.filename or f"uploaded_{index}.txt").name
        target_path = UPLOAD_CASE_DIR / safe_name

        content_bytes = file.file.read()
        target_path.write_bytes(content_bytes)

        extracted_text = extract_text_for_generation(target_path)
        doc_type, label = classify_document(safe_name, extracted_text)

        documents.append(
            {
                "name": safe_name,
                "type": doc_type,
                "label": label,
                "exists": True,
                "extractable": bool(extracted_text),
                "type_confirmed": False,
                "auto_type": doc_type,
            }
        )

    return documents


def generate_outputs_for_uploaded_case():
    global STATE

    if STATE["mode"] != "uploaded":
        raise HTTPException(status_code=400, detail="当前不是导入案例，无法开始解析。")

    documents = STATE["documents"]
    if not documents:
        raise HTTPException(status_code=400, detail="当前没有已导入文件。")
    capabilities = derive_input_capabilities(documents)
    if capabilities.get("unconfirmed_type_count", 0) > 0:
        raise HTTPException(status_code=400, detail="请先确认全部材料类型，再开始解析。")
    if not capabilities.get("can_generate_resume"):
        raise HTTPException(status_code=400, detail="当前材料不足以解析。请至少导入并确认一份简历。")

    text_map = collect_uploaded_texts(documents)
    ok, vmsg = validate_materials_for_parse(documents, text_map)
    if not ok:
        raise HTTPException(status_code=400, detail=vmsg)

    backend = os.getenv("LLM_BACKEND", "ollama").lower().strip()
    if backend in ("ollama", ""):
        if not check_ollama_available():
            raise HTTPException(
                status_code=503,
                detail="本地 Ollama 不可用：请确认已在 http://localhost:11434 启动并拉取模型后再生成。",
            )

    clear_directory(GENERATED_OUTPUT_DIR)
    catalog = build_snippet_catalog(documents, text_map)
    cmap = catalog_to_map(catalog)
    valid_ids = set(cmap.keys())

    ctx = build_import_context(
        documents,
        text_map,
        catalog,
        goal_mode=STATE.get("goal_mode", "both"),
    )

    gm = STATE.get("goal_mode", "both")
    try:
        resume_raw = (
            _run_llm_json_pack(messages_resume_delta(ctx), ResumeDeltaPayload)
            if capabilities.get("can_generate_resume_suggestions")
            else ResumeDeltaPayload()
        )
        if gm == "interview":
            job_raw = JobBriefPayload()
            interview_raw = (
                _run_llm_json_pack(messages_interview_pack(ctx), InterviewPackPayload)
                if capabilities.get("can_generate_interview_prep")
                else InterviewPackPayload()
            )
        elif gm == "delivery":
            job_raw = (
                _run_llm_json_pack(messages_job_brief(ctx), JobBriefPayload)
                if capabilities.get("can_generate_job_brief")
                else JobBriefPayload()
            )
            interview_raw = InterviewPackPayload()
        else:
            job_raw = (
                _run_llm_json_pack(messages_job_brief(ctx), JobBriefPayload)
                if capabilities.get("can_generate_job_brief")
                else JobBriefPayload()
            )
            interview_raw = (
                _run_llm_json_pack(messages_interview_pack(ctx), InterviewPackPayload)
                if capabilities.get("can_generate_interview_prep")
                else InterviewPackPayload()
            )
    except ValueError as exc:
        raise HTTPException(status_code=502, detail=f"模型输出无法解析为 JSON：{exc}") from exc
    except ValidationError as exc:
        raise HTTPException(status_code=502, detail=f"JSON 与约定结构不一致：{exc}") from exc
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"调用语言模型失败：{exc}") from exc

    job_payload = _sanitize_payload(job_raw, valid_ids, catalog)
    resume_payload = _sanitize_payload(resume_raw, valid_ids, catalog)
    interview_payload = _sanitize_payload(interview_raw, valid_ids, catalog)
    job_payload, resume_payload, interview_payload = apply_goal_mode_truncation(
        gm, job_payload, resume_payload, interview_payload
    )

    now = datetime.now().isoformat(timespec="seconds")
    model_name = get_active_model_label()
    gen_meta = _generation_meta_block(model_name, now)
    missing = compute_missing_inputs(STATE.get("goal_mode", "both"), documents)
    acceptance = compute_acceptance_hints(documents, catalog, STATE.get("goal_mode", "both"), missing)

    archive_manifest = {
        "case_id": STATE["case_id"],
        "case_title": STATE["title"],
        "generated_at": now,
        "input_files": [
            {
                "name": doc["name"],
                "type": doc["type"],
                "label": doc["label"],
                "extractable": doc["extractable"],
            }
            for doc in documents
        ],
        "generated_outputs": [OUTPUT_FILE_MAP[i][1] for i in output_ids_for_capabilities(capabilities, include_manifest=True)],
        "notes": "解析摘要由确定性逻辑生成；正文由 LLM 依据导入材料与 snippet 候选生成 JSON 后渲染。",
    }

    write_generated_outputs(GENERATED_OUTPUT_DIR, job_payload, resume_payload, interview_payload, archive_manifest)

    output_definitions = build_output_definitions(
        GENERATED_OUTPUT_DIR,
        cmap,
        catalog,
        job_payload,
        resume_payload,
        interview_payload,
        gen_meta,
        acceptance,
        capabilities,
    )

    case_id = STATE["case_id"]
    rt = case_runtime_dir(ROOT_DIR, case_id)
    prev_ex = load_execution_state(rt / "execution.json") or {}
    prev_gen = prev_ex.get("generation_id")
    prev_files = list(prev_ex.get("input_files", []))
    prev_missing = list(prev_ex.get("missing_inputs", []))
    prev_suggestions = list(prev_ex.get("suggestion_items", []))

    gen_id = new_generation_id()
    new_files = [d["name"] for d in documents]
    new_items = resume_payload_to_suggestion_items(resume_payload, gen_id)
    merged_items = merge_suggestion_status(prev_suggestions, new_items)
    comparison = build_comparison_summary(
        generation_id=gen_id,
        previous_generation_id=prev_gen,
        prev_files=prev_files,
        new_files=new_files,
        prev_missing=prev_missing,
        new_missing=missing,
        prev_suggestions=prev_suggestions,
        new_suggestions=merged_items,
    )

    resume_body = collect_typed_text(documents, text_map, "resume")
    interview_body = collect_typed_text(documents, text_map, "interview_note")
    rp, ip = working_draft_paths(rt)
    ensure_structured_resume_draft(rp, resume_body)
    ensure_interview_talking_draft(ip, interview_body)
    snap_md = read_text_file(rp)
    save_original_resume_snapshot(rt, snap_md)
    if capabilities.get("can_generate_interview_prep"):
        write_interview_expression_coaching(rt, documents, text_map)

    save_execution_state(
        rt / "execution.json",
        {
            "generation_id": gen_id,
            "previous_generation_id": prev_gen,
            "input_files": new_files,
            "missing_inputs": missing,
            "suggestion_items": merged_items,
            "comparison_summary": comparison,
            "last_updated": now,
            "workflow_phase": "parsed",
            "input_capabilities": capabilities,
            "input_snapshot_at_generation": build_upload_snapshot(UPLOAD_CASE_DIR, new_files),
        },
    )
    save_draft_apply_summary(rt, merged_items)

    STATE["outputs"] = build_output_list("ready", capabilities)
    STATE["output_definitions"] = output_definitions
    STATE["last_generated_at"] = now
    STATE["acceptance_hints"] = acceptance
    STATE["input_capabilities"] = capabilities
    STATE["workflow_phase"] = "parsed"
    STATE["execution"] = enrich_execution(case_id)
    refresh_derived_case_fields()

    return {"ok": True, "message": "解析完成：建议与材料依据已就绪，可继续在工作稿中确认与写入。", "comparison": comparison}


@app.get("/health")
def health():
    payload = {"ok": True, "message": "backend is running"}
    if os.getenv("LLM_BACKEND", "ollama").lower().strip() in ("ollama", ""):
        payload["ollama_reachable"] = check_ollama_available()
    return payload


@app.get("/api/current-case")
def current_case():
    state_without_defs = {k: v for k, v in STATE.items() if k != "output_definitions"}
    
    # 如果是空白态，返回基本结构但不包含执行数据
    if STATE.get("mode") == "blank":
        state_without_defs["execution"] = None
        return state_without_defs
    
    # 否则处理正常的案例（demo 或 uploaded）
    cid = state_without_defs.get("case_id")
    if cid:
        ex = enrich_execution(cid)
        upload_root = INPUT_DIR if state_without_defs.get("mode") == "demo" else UPLOAD_CASE_DIR
        ex["readiness_summary"] = build_readiness_summary(
            documents=state_without_defs.get("documents") or [],
            missing_inputs=state_without_defs.get("missing_inputs") or [],
            acceptance_hints=state_without_defs.get("acceptance_hints"),
            execution=ex,
            upload_root=upload_root,
            has_generated=bool(ex.get("generation_id")),
        )
        state_without_defs["execution"] = ex
    return state_without_defs


@app.get("/api/ai-status")
def api_ai_status():
    backend = os.getenv("LLM_BACKEND", "ollama").lower().strip()
    if backend in ("ollama", ""):
        available = check_ollama_available()
    else:
        available = True
    return {
        "available": available,
        "backend": backend,
        "model": get_active_model_label(),
        "affected_when_unavailable": [
            "材料解析（简历建议；材料足够时生成岗位理解和面试准备）",
            "建议「写入工作稿」时的正文改写",
            "面试表达辅导正文生成",
            "面试练习题与参考回答生成",
            "逐题评分与综合报告",
        ],
    }


@app.get("/api/document-preview/{filename}")
def api_document_preview(filename: str):
    safe = Path(filename).name
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    docs = STATE.get("documents") or []
    if not any(d.get("name") == safe for d in docs):
        raise HTTPException(status_code=404, detail="文件不在当前案例材料列表中。")
    root = INPUT_DIR if STATE.get("mode") == "demo" else UPLOAD_CASE_DIR
    p = root / safe
    if not p.is_file():
        raise HTTPException(status_code=404, detail="文件不存在。")
    return {"name": safe, "content": extract_text_for_generation(p)[:120000]}


@app.post("/api/hr-persona")
def api_hr_persona(body: HrPersonaBody):
    global STATE
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    rt = case_runtime_dir(ROOT_DIR, cid)
    cur = load_hr_persona(rt)
    if cur.get("locked"):
        raise HTTPException(status_code=400, detail="本轮已锁定 HR 风格，不可更改。")
    pid = (body.persona_id or "").strip() or "calm_rational"
    data = {
        "persona_id": pid,
        "locked": True,
        "label": (body.label or "").strip() or pid,
    }
    rt.mkdir(parents=True, exist_ok=True)
    hr_persona_path(rt).write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    documents = STATE.get("documents") or []
    if STATE.get("mode") == "demo":
        text_map = {
            d["name"]: extract_text_for_generation(INPUT_DIR / d["name"])
            for d in documents
            if (INPUT_DIR / d["name"]).is_file()
        }
    else:
        text_map = collect_uploaded_texts(documents)
    try:
        write_interview_expression_coaching(rt, documents, text_map)
    except Exception:
        pass
    STATE["execution"] = enrich_execution(cid)
    return {"ok": True, "hr_persona": data}


@app.get("/api/current-output/{output_id}")
def current_output(output_id: str):
    definition = STATE["output_definitions"].get(output_id)
    if not definition:
        raise HTTPException(status_code=404, detail="当前解析结果尚未生成，请先点击「开始解析」。")

    file_path = Path(definition["file_path"])
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="结果文件不存在。")

    raw_content = read_text_file(file_path)

    try:
        if definition["format"] == "json":
            content = json.dumps(json.loads(raw_content), ensure_ascii=False, indent=2)
        else:
            content = raw_content
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"结果文件解析失败：{str(exc)}")

    return {
        "output_id": output_id,
        "title": definition["title"],
        "format": definition["format"],
        "content": content,
        "evidence": definition["evidence"],
        "generation_meta": definition.get("generation_meta"),
        "acceptance": definition.get("acceptance"),
    }


@app.patch("/api/suggestions/{item_id}")
def patch_suggestion_status(item_id: str, body: SuggestionStatusPatch):
    global STATE
    if body.status != "accepted":
        raise HTTPException(status_code=400, detail="仅支持接受操作；要忽略请使用「忽略」删除该条。")

    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")

    rt = case_runtime_dir(ROOT_DIR, cid)
    path = rt / "execution.json"
    ex = load_execution_state(path)
    if not ex:
        raise HTTPException(status_code=404, detail="尚无执行数据，请先完成「开始解析」。")

    items = ex.get("suggestion_items", [])
    found = None
    for it in items:
        if it.get("id") == item_id:
            found = it
            break
    if not found:
        raise HTTPException(status_code=404, detail="未找到该建议。")

    if found.get("applied_to_draft"):
        raise HTTPException(status_code=400, detail="该建议已写入工作稿。")

    if found.get("status") == "accepted":
        save_execution_state(path, ex)
        STATE["execution"] = enrich_execution(cid)
        return {"ok": True}

    found["status"] = "accepted"

    save_execution_state(path, ex)
    save_draft_apply_summary(rt, items)
    STATE["execution"] = enrich_execution(cid)
    return {"ok": True}


@app.delete("/api/suggestions/{item_id}")
def delete_suggestion(item_id: str):
    """忽略：撤销该建议对正文的影响，并从列表移除。"""
    global STATE
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")

    rt = case_runtime_dir(ROOT_DIR, cid)
    path = rt / "execution.json"
    ex = load_execution_state(path)
    if not ex:
        raise HTTPException(status_code=404, detail="尚无执行数据。")

    items = ex.get("suggestion_items", [])
    found = None
    for it in items:
        if it.get("id") == item_id:
            found = it
            break
    if not found:
        raise HTTPException(status_code=404, detail="未找到该建议。")

    rp, _ = working_draft_paths(rt)
    if found.get("applied_to_draft") and found.get("pre_apply_section_body") is not None:
        try:
            replace_section_body_in_file(rp, found.get("target_section") or "general", found["pre_apply_section_body"])
        except Exception:
            pass
    elif found.get("status") == "accepted" and not found.get("applied_to_draft"):
        raw = read_text_file(rp)
        raw2 = remove_improvement_block(raw, item_id)
        if raw2 != raw:
            rp.write_text(raw2, encoding="utf-8")

    new_items = [it for it in items if it.get("id") != item_id]
    ex["suggestion_items"] = new_items
    save_execution_state(path, ex)
    save_draft_apply_summary(rt, new_items)
    STATE["execution"] = enrich_execution(cid)
    return {"ok": True}


@app.post("/api/suggestions/{item_id}/apply-to-draft")
def apply_suggestion_to_draft(item_id: str):
    global STATE
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")

    rt = case_runtime_dir(ROOT_DIR, cid)
    path = rt / "execution.json"
    ex = load_execution_state(path)
    if not ex:
        raise HTTPException(status_code=404, detail="尚无执行数据，请先完成「开始解析」。")

    items = ex.get("suggestion_items", [])
    found = None
    for it in items:
        if it.get("id") == item_id:
            found = it
            break
    if not found:
        raise HTTPException(status_code=404, detail="未找到该建议。")

    if found.get("applied_to_draft"):
        raise HTTPException(status_code=400, detail="该建议已写入工作稿。")
    if found.get("status") != "accepted":
        found["status"] = "accepted"

    backend = os.getenv("LLM_BACKEND", "ollama").lower().strip()
    if backend in ("ollama", "") and not check_ollama_available():
        raise HTTPException(status_code=503, detail="模型不可用，无法根据材料改写正文。")

    documents = STATE.get("documents") or []
    if STATE.get("mode") == "demo":
        text_map: Dict[str, str] = {}
        for d in documents:
            fp = INPUT_DIR / d["name"]
            text_map[d["name"]] = extract_text_for_generation(fp) if fp.exists() else ""
    else:
        text_map = collect_uploaded_texts(documents)

    blob = build_materials_blob(documents, text_map)
    rp, _ = working_draft_paths(rt)
    raw = read_text_file(rp)
    key = normalize_target_section(found.get("target_section") or "general")
    body_full = get_section_body_text(raw, key)
    found["pre_apply_section_body"] = body_full

    clean = strip_improvement_blocks_for_llm(body_full)
    nk = key if key != "general" else "summary"
    section_title = SECTION_TITLES.get(nk, "简历")

    try:
        new_body = rewrite_resume_section_with_llm(
            section_title=section_title,
            current_section_body=clean,
            suggestion_text=(found.get("text") or "").strip(),
            materials_blob=blob,
        )
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"改写失败：{exc}") from exc

    replace_section_body_in_file(rp, found.get("target_section") or "general", new_body)

    applied_at = datetime.now().isoformat(timespec="seconds")
    found["applied_to_draft"] = True
    found["applied_at"] = applied_at
    save_execution_state(path, ex)
    save_draft_apply_summary(rt, items)
    STATE["execution"] = enrich_execution(cid)
    return {"ok": True, "applied_at": applied_at}


@app.get("/api/original-resume-snapshot")
def get_original_resume_snapshot():
    """解析后冻结的简历各区块原文，用于「对比原文」。"""
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    rt = case_runtime_dir(ROOT_DIR, cid)
    return {"sections": load_original_resume_snapshot(rt)}


@app.get("/api/working-drafts/{which}")
def get_working_draft(which: str):
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    if which not in ("resume", "interview"):
        raise HTTPException(status_code=400, detail="which 须为 resume | interview")

    rt = case_runtime_dir(ROOT_DIR, cid)
    rp, ip = working_draft_paths(rt)
    target = rp if which == "resume" else ip
    if not target.exists():
        return {"content": "", "exists": False, "sections": []}
    raw = read_text_file(target)
    if which == "resume" and "### 原始内容" in raw:
        raw = migrate_old_template_to_clean(raw)
        write_text_file(target, raw)
    out: Dict[str, Any] = {"content": raw, "exists": True}
    if which == "resume":
        ex = load_execution_state(rt / "execution.json") or {}
        items = ex.get("suggestion_items") or []
        orig = load_original_resume_snapshot(rt)
        out["sections"] = parse_resume_draft_sections(raw)
        out["resume_sections"] = build_resume_section_views(raw, items, orig)
        out["pending_accept_overlays"] = build_pending_accept_overlays(items)
    else:
        out["sections"] = parse_resume_draft_sections(raw)
        out["resume_sections"] = []
    return out


@app.put("/api/working-drafts/{which}")
def put_working_draft(which: str, body: WorkingDraftPut):
    global STATE
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    if which not in ("resume", "interview"):
        raise HTTPException(status_code=400, detail="which 须为 resume | interview")

    rt = case_runtime_dir(ROOT_DIR, cid)
    rp, ip = working_draft_paths(rt)
    target = rp if which == "resume" else ip
    ensure_dir(target.parent)
    write_text_file(target, body.content)
    STATE["execution"] = enrich_execution(cid)
    return {"ok": True}


@app.post("/api/refined-draft/generate")
def generate_refined_draft():
    """生成独立润色稿（写入 refined/，不修改 resume_working_draft.md）。"""
    global STATE
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    rt = case_runtime_dir(ROOT_DIR, cid)
    rp, _ = working_draft_paths(rt)
    working_md = read_text_file(rp)
    if len(working_md.strip()) < 30:
        raise HTTPException(
            status_code=400,
            detail="简历工作稿过短或不存在，请先在工作稿中保留有效内容。",
        )
    ex = load_execution_state(rt / "execution.json") or {}
    gen_id = ex.get("generation_id")
    items = ex.get("suggestion_items") or []
    gm = STATE.get("goal_mode", "both")
    missing = STATE.get("missing_inputs") or []
    sug_summary = build_accepted_applied_summary(items)
    messages = messages_refined_resume(
        working_draft_markdown=working_md,
        goal_mode=gm,
        missing_inputs=missing,
        suggestion_summary=sug_summary,
    )
    try:
        payload = _run_llm_json_pack(messages, RefinedResumePayload)
    except ValueError as exc:
        raise HTTPException(status_code=502, detail=f"模型输出无法解析为 JSON：{exc}") from exc
    except ValidationError as exc:
        raise HTTPException(status_code=502, detail=f"JSON 与约定结构不一致：{exc}") from exc
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"调用语言模型失败：{exc}") from exc

    refined_id = new_refined_id()
    model_name = get_active_model_label()
    _, _, meta = save_refined_artifacts(
        rt,
        refined_id,
        payload,
        based_on_generation_id=gen_id,
        working_draft_sha256=hash_text(working_md),
        model_name=model_name,
        goal_mode=gm,
    )
    STATE["execution"] = enrich_execution(cid)
    return {"ok": True, "refined_id": refined_id, "meta": meta}


@app.get("/api/refined-draft")
def get_refined_draft_endpoint(refined_id: Optional[str] = Query(None)):
    """获取指定或最新润色稿全文与元信息。"""
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    rt = case_runtime_dir(ROOT_DIR, cid)
    if refined_id:
        data = load_refined_by_id(rt, refined_id)
    else:
        data = load_latest_refined(rt)
    if not data:
        raise HTTPException(status_code=404, detail="尚无润色稿，请先在右侧生成。")
    return data


@app.get("/api/refined-draft/compare")
def get_refined_draft_compare():
    """当前工作稿 vs 最新润色稿的对照摘要（规则生成，非 LLM）。"""
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    rt = case_runtime_dir(ROOT_DIR, cid)
    rp, _ = working_draft_paths(rt)
    working_md = read_text_file(rp)
    ref = load_latest_refined(rt)
    if not ref:
        raise HTTPException(status_code=404, detail="尚无润色稿。")
    meta = ref.get("meta") or {}
    body = ref.get("markdown_body") or ""
    return build_compare_summary(
        working_draft_markdown=working_md,
        refined_markdown_body=body,
        meta=meta,
    )


@app.get("/api/export/deliverables-markdown")
def export_deliverables_markdown():
    """合并当前解析后 Markdown，供用户另存为本地文件。"""
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    caps = STATE.get("input_capabilities") or derive_input_capabilities(STATE.get("documents") or [])
    allowed_ids = set(output_ids_for_capabilities(caps, include_manifest=False))
    chunks: List[str] = []
    for output_id in ("job_brief", "resume_delta", "interview_pack"):
        if output_id not in allowed_ids:
            continue
        fname = OUTPUT_FILE_MAP[output_id][1]
        p = GENERATED_OUTPUT_DIR / fname
        if p.is_file():
            chunks.append(f"# {fname}\n\n{read_text_file(p)}")
    return {
        "ok": True,
        "markdown": "\n\n---\n\n".join(chunks) if chunks else "",
        "suggested_filename": "解析后内容汇总.md",
    }


@app.get("/api/interview-practice")
def get_interview_practice_data():
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    rt = case_runtime_dir(ROOT_DIR, cid)
    data = load_practice_pack(rt)
    if not data:
        raise HTTPException(status_code=404, detail="尚未生成面试练习题。")
    return data


@app.post("/api/interview-practice/generate")
def post_interview_practice_generate():
    global STATE
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    documents = STATE.get("documents") or []
    if not documents:
        raise HTTPException(status_code=400, detail="请先导入材料。")
    caps = derive_input_capabilities(documents)
    if not caps.get("can_generate_practice"):
        raise HTTPException(status_code=400, detail="当前输入不足以支撑面试练习题。请至少提供简历、岗位说明，以及面试记录或项目补充材料。")
    backend = os.getenv("LLM_BACKEND", "ollama").lower().strip()
    if backend in ("ollama", ""):
        if not check_ollama_available():
            raise HTTPException(
                status_code=503,
                detail="本地 Ollama 不可用：请确认已在 http://localhost:11434 启动并拉取模型。",
            )
    if STATE.get("mode") == "demo":
        text_map = {
            d["name"]: extract_text_for_generation(INPUT_DIR / d["name"])
            for d in documents
            if (INPUT_DIR / d["name"]).is_file()
        }
    else:
        text_map = collect_uploaded_texts(documents)
    rt = case_runtime_dir(ROOT_DIR, cid)
    payload = generate_interview_practice_pack_safe(documents=documents, text_map=text_map)
    save_practice_pack(rt, payload, answers={}, full_report=None)
    return {"ok": True, **(load_practice_pack(rt) or {})}


class InterviewPracticeSubmitBody(BaseModel):
    answers: Dict[str, str] = Field(default_factory=dict)


@app.post("/api/interview-practice/submit-all")
def post_interview_practice_submit_all(body: InterviewPracticeSubmitBody):
    global STATE
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    backend = os.getenv("LLM_BACKEND", "ollama").lower().strip()
    if backend in ("ollama", ""):
        if not check_ollama_available():
            raise HTTPException(
                status_code=503,
                detail="本地 Ollama 不可用：综合分析报告需要模型生成。",
            )
    documents = STATE.get("documents") or []
    caps = derive_input_capabilities(documents)
    if not caps.get("can_generate_report"):
        raise HTTPException(status_code=400, detail="当前输入不足以支撑完整综合分析报告。请补充岗位说明和面试记录/项目材料后再试。")
    if STATE.get("mode") == "demo":
        text_map = {
            d["name"]: extract_text_for_generation(INPUT_DIR / d["name"])
            for d in documents
            if (INPUT_DIR / d["name"]).is_file()
        }
    else:
        text_map = collect_uploaded_texts(documents)
    rt = case_runtime_dir(ROOT_DIR, cid)
    pack = load_practice_pack(rt)
    if not pack or not pack.get("questions"):
        raise HTTPException(status_code=400, detail="请先生成面试练习题。")
    qs = pack["questions"]
    ans = body.answers or {}
    for q in qs:
        qid = q.get("id")
        if not (ans.get(qid) or "").strip():
            raise HTTPException(status_code=400, detail="未全部作答，无法生成综合分析报告")
        if len((ans.get(qid) or "").strip()) < 20:
            raise HTTPException(status_code=400, detail="报告暂未生成成功，请补全答案后重试。")
    try:
        report = submit_practice_full_report_safe(
            documents=documents,
            text_map=text_map,
            questions=qs,
            answers=ans,
        )
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc) or "报告暂未生成成功，请补全答案后重试。") from exc
    payload = InterviewPracticePayload.model_validate({"questions": qs})
    save_practice_pack(
        rt,
        payload,
        answers=ans,
        full_report=report.model_dump(),
    )
    return {"ok": True, "full_report": report.model_dump()}


@app.get("/api/interview-expression")
def get_interview_expression_coaching():
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    rt = case_runtime_dir(ROOT_DIR, cid)
    p = rt / "interview_expression_coaching.md"
    if not p.is_file():
        return {"content": "", "exists": False}
    return {"content": read_text_file(p), "exists": True}


@app.post("/api/export/result-pack-zip")
def export_result_pack_zip(body: ResultPackExportBody):
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    ex = load_execution_state(case_runtime_dir(ROOT_DIR, cid) / "execution.json") or {}
    if ex.get("workflow_phase") != "parsed" and not ex.get("generation_id"):
        raise HTTPException(status_code=400, detail="请先完成「开始解析」并确认材料已生成。")
    rt = case_runtime_dir(ROOT_DIR, cid)
    rp, _ = working_draft_paths(rt)
    resume_md = read_text_file(rp)
    job_brief_md = read_generated_md(GENERATED_OUTPUT_DIR, "job_brief.md")
    resume_delta_md = read_generated_md(GENERATED_OUTPUT_DIR, "resume_delta.md")
    interview_pack_md = read_generated_md(GENERATED_OUTPUT_DIR, "interview_pack.md")
    report_md = ""
    if body.include_interview_report:
        pp = load_practice_pack(rt)
        if pp and pp.get("full_report"):
            fr = pp["full_report"]
            report_md = (
                f"# 面试练习综合分析\n\n## 总体评价\n{fr.get('overall','')}\n\n"
                f"## 岗位匹配\n{fr.get('role_fit_judgement','')}\n\n"
                f"## 风险与亮点\n{fr.get('hire_risk_and_highlights','')}\n\n"
                f"## 补强点\n{fr.get('top_gaps','')}\n"
            )
    zip_bytes = build_result_pack_zip_bytes(
        resume_markdown=resume_md,
        job_brief_md=job_brief_md,
        resume_delta_md=resume_delta_md,
        interview_pack_md=interview_pack_md,
        interview_report_md=report_md or None,
        include_interview_report=body.include_interview_report and bool(report_md.strip()),
        capabilities=STATE.get("input_capabilities") or derive_input_capabilities(STATE.get("documents") or []),
    )
    base = sanitize_zip_base(STATE.get("title") or "求职材料")
    ascii_name = base.encode("ascii", "ignore").decode("ascii") or "case"
    fn_utf8 = f"{base}_解析结果.zip"
    disp = f'attachment; filename="{ascii_name}_result.zip"; filename*=UTF-8\'\'{quote(fn_utf8)}'
    return Response(
        content=zip_bytes,
        media_type="application/zip",
        headers={"Content-Disposition": disp},
    )


@app.post("/api/parse-materials")
def post_parse_materials():
    """阶段 A：开始解析（校验材料 + 生成建议与工作稿依据）。"""
    return generate_outputs_for_uploaded_case()


@app.post("/api/interview-practice/score")
def post_interview_practice_score(body: InterviewScoreBody):
    backend = os.getenv("LLM_BACKEND", "ollama").lower().strip()
    if backend in ("ollama", ""):
        if not check_ollama_available():
            raise HTTPException(
                status_code=503,
                detail="本地 Ollama 不可用：请确认已在 http://localhost:11434 启动并拉取模型。",
            )
    cid = STATE.get("case_id")
    if not cid:
        raise HTTPException(status_code=400, detail="无当前案例。")
    rt = case_runtime_dir(ROOT_DIR, cid)
    persona = load_hr_persona(rt)
    try:
        r = score_interview_answer(
            question=body.question,
            user_answer=body.user_answer,
            reference_answer=body.reference_answer or "",
            persona_id=str(persona.get("persona_id") or "calm_rational"),
        )
        return r.model_dump()
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"评分失败：{exc}") from exc


@app.post("/api/cases/new")
def create_new_case():
    global STATE
    clear_directory(UPLOAD_CASE_DIR)
    STATE = build_uploaded_state([], title="未命名案例", goal_mode="both")
    refresh_derived_case_fields()
    return {"ok": True, "message": "已新建案例，可导入材料。"}


@app.patch("/api/case")
def patch_case(body: CasePatch):
    global STATE
    if STATE.get("mode") == "demo":
        raise HTTPException(status_code=400, detail="示例案例下不可修改标题与模式。")
    if body.title is not None:
        STATE["title"] = body.title.strip() or "未命名案例"
    if body.goal_mode is not None:
        if body.goal_mode not in ("delivery", "interview", "both"):
            raise HTTPException(status_code=400, detail="goal_mode 须为 delivery | interview | both")
        STATE["goal_mode"] = body.goal_mode
    refresh_derived_case_fields()
    return {"ok": True}


@app.patch("/api/documents/{filename}/type")
def patch_document_type(filename: str, body: DocTypePatch):
    global STATE
    if STATE.get("mode") == "demo":
        raise HTTPException(status_code=400, detail="示例案例下不可修改类型。")
    safe = Path(filename).name
    if body.doc_type not in DOC_LABELS:
        raise HTTPException(status_code=400, detail="不支持的文档类型。")
    for doc in STATE.get("documents", []):
        if doc["name"] == safe:
            doc["type"] = body.doc_type
            doc["label"] = DOC_LABELS[body.doc_type]
            refresh_derived_case_fields()
            return {"ok": True}
    raise HTTPException(status_code=404, detail="未找到该文件。")


@app.post("/api/documents/{filename}/confirm-type")
def confirm_document_type(filename: str):
    global STATE
    if STATE.get("mode") == "demo":
        raise HTTPException(status_code=400, detail="示例案例无需确认。")
    safe = Path(filename).name
    for doc in STATE.get("documents", []):
        if doc["name"] == safe:
            doc["type_confirmed"] = True
            refresh_derived_case_fields()
            return {"ok": True}
    raise HTTPException(status_code=404, detail="未找到该文件。")


@app.delete("/api/documents/{filename}")
def delete_document(filename: str):
    """Remove one uploaded file from the current case (disk + state)."""
    global STATE
    if STATE.get("mode") == "demo":
        raise HTTPException(status_code=400, detail="示例案例不可删除材料。")
    if STATE.get("mode") != "uploaded":
        raise HTTPException(status_code=400, detail="当前状态不可删除材料。")
    safe = Path(filename).name
    docs = STATE.get("documents") or []
    new_docs = [d for d in docs if d.get("name") != safe]
    if len(new_docs) == len(docs):
        raise HTTPException(status_code=404, detail="未找到该文件。")
    target = UPLOAD_CASE_DIR / safe
    try:
        if target.is_file():
            target.unlink()
    except OSError:
        pass
    STATE["documents"] = new_docs
    st = STATE.get("stats") or {}
    st["document_count"] = len(new_docs)
    STATE["stats"] = st
    refresh_derived_case_fields()
    return {"ok": True, "message": "已移除该材料。"}


@app.post("/api/upload-documents")
async def upload_documents(files: List[UploadFile] = File(...)):
    global STATE

    if not files:
        raise HTTPException(status_code=400, detail="没有收到文件。")

    prev = STATE if STATE.get("mode") == "uploaded" else None
    cid = STATE.get("case_id") if STATE.get("mode") == "uploaded" else None
    goal = STATE.get("goal_mode", "both") if STATE.get("mode") == "uploaded" else "both"

    documents = save_uploaded_files(files)
    first_stem = Path(documents[0]["name"]).stem if documents else "未命名案例"
    old_title = (prev or {}).get("title") or ""
    defaults = {"", "未命名案例", "当前求职案例"}
    if prev and old_title.strip() not in defaults:
        title = old_title.strip()
    else:
        title = first_stem

    STATE = build_uploaded_state(documents, case_id=cid, title=title, goal_mode=goal)
    refresh_derived_case_fields()

    return {
        "ok": True,
        "message": f"已导入 {len(documents)} 个文件。",
        "documents": documents,
    }


@app.post("/api/generate-output-pack")
def generate_output_pack():
    return generate_outputs_for_uploaded_case()


@app.post("/api/reset-demo")
def reset_demo():
    global STATE
    rt = case_runtime_dir(ROOT_DIR, "case_demo_001")
    hp = hr_persona_path(rt)
    if hp.is_file():
        try:
            hp.unlink()
        except OSError:
            pass
    STATE = build_demo_state()
    return {"ok": True, "message": "已恢复示例案例。"}
